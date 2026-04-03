import os
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Инициализация объекта Document
doc = docx.Document()

# Настройка глобальных параметров шрифта по умолчанию
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

# Настройка полей страницы (стандартные отступы)
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)

def add_paragraph(doc_obj, text, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, bold=False, size=14, spacing=1.5):
    # Добавление абзаца с параметризированным форматированием
    p = doc_obj.add_paragraph()
    run = p.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.line_spacing = spacing
    p.alignment = align
    return p

def read_file_content(filepath):
    # Чтение данных из исходного файла (UTF-8)
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as file:
            return file.read()
    return f"Файл не найден: {filepath}"

def add_code_block(doc_obj, code_text):
    # Вставка блока кода с применением моноширинного шрифта
    p = doc_obj.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# --- ГЕНЕРАЦИЯ ТИТУЛЬНОГО ЛИСТА ---
add_paragraph(doc, 'Министерство науки и высшего образования Российской Федерации', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
add_paragraph(doc, 'Муромский институт (филиал)', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
add_paragraph(doc, 'федерального государственного бюджетного образовательного учреждения высшего\nобразования', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
add_paragraph(doc, '«Владимирский государственный университет\nимени Александра Григорьевича и Николая Григорьевича Столетовых»', align=WD_PARAGRAPH_ALIGNMENT.CENTER)

add_paragraph(doc, '\nФакультет: ИТР', align=WD_PARAGRAPH_ALIGNMENT.LEFT)
add_paragraph(doc, 'Кафедра: ПИН\n', align=WD_PARAGRAPH_ALIGNMENT.LEFT)

add_paragraph(doc, 'ЛАБОРАТОРНАЯ РАБОТА №1', align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True, size=16)
add_paragraph(doc, 'По: Модульному тестированию', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
add_paragraph(doc, 'Тема: RomanNumeral Converter - арабские римские числа', align=WD_PARAGRAPH_ALIGNMENT.CENTER)

# Блок с подписями (форматирование пробелами для эмуляции позиционирования)
add_paragraph(doc, '\n\n\nРуководитель\t\t\t\t\t\t\tФамилия И.О.', align=WD_PARAGRAPH_ALIGNMENT.LEFT)
add_paragraph(doc, '\t\t\t\t(подпись)\t\t(дата)', align=WD_PARAGRAPH_ALIGNMENT.LEFT, size=10)
add_paragraph(doc, '\nСтудент группы ПИН-xxx\t\t\t\t\tФамилия И.О.', align=WD_PARAGRAPH_ALIGNMENT.LEFT)
add_paragraph(doc, '\t\t\t\t(подпись)\t\t(дата)', align=WD_PARAGRAPH_ALIGNMENT.LEFT, size=10)

add_paragraph(doc, '\n\n\nМуром 2026', align=WD_PARAGRAPH_ALIGNMENT.CENTER)

# Добавление разрыва страницы
doc.add_page_break()

# --- ОСНОВНОЙ КОНТЕНТ ---
add_paragraph(doc, 'Лабораторная работа №1', align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
add_paragraph(doc, 'Тема: RomanNumeral Converter - арабские римские числа', bold=True)
add_paragraph(doc, 'Цель работы: Освоить базовые навыки применения модульного тестирования с акцентом на выявление дефектов и оценку качества реализации программных компонентов.\n')

add_paragraph(doc, '1. Техническое задание и промпт', align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
add_paragraph(doc, 'Промпт: Сгенерируй класс RomanNumeralConverter на Java со следующими методами: 1. Метод: toRoman, входные параметры: int number, возвращает: String. 2. Метод: toArabic, входные параметры: String roman, возвращает: int. Особые условия: метод toRoman должен выбрасывать IllegalArgumentException при значениях меньше 1 или больше 3999. Метод toArabic должен выбрасывать IllegalArgumentException при передаче пустой строки или некорректного формата римского числа. Класс должен реализовывать интерфейс IRomanNumeralConverter, находящийся в пакете lab.interfaces.')

add_paragraph(doc, '2. Исходный код', align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
add_paragraph(doc, 'Реализации сгенерированы моделями Kimi K2, Model2 и Model3. Ниже представлены листинги исходного кода.')

# Интеграция файлов исходного кода
models = ['gencode1', 'gencode2', 'gencode3']
for model in models:
    add_paragraph(doc, f'Реализация от {model}:', bold=True)
    path = f"src/lab/implementations/{model}/RomanNumeralConverter.java"
    code = read_file_content(path)
    add_code_block(doc, code)

add_paragraph(doc, 'Анализ различий: Код от первой модели включает валидацию каноничной формы числа. Вторая модель не осуществляет приведение строки к верхнему регистру. Третья модель предоставляет базовую работоспособную логику.')

doc.add_page_break()

add_paragraph(doc, '3. Тесты', align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
add_paragraph(doc, 'Ниже представлен исходный код модульных тестов.')

test_path = "src/lab/tests/GenCode1Tests.java"
test_code = read_file_content(test_path)
add_code_block(doc, test_code)

add_paragraph(doc, 'Анализ падения тестов: Тесты GenCode2 и GenCode3 упали на проверке текста исключений ("Invalid Roman numeral" vs "Некорректный формат..."), так как модели сгенерировали сообщения на разных языках, а проверка assertThrows ожидала строгое совпадение строк. Также GenCode2 не прошел тест на обработку нижнего регистра (toArabic_LowerCaseString_ProcessedSuccessfully).')

add_paragraph(doc, '4. Выводы', align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
add_paragraph(doc, 'Оценка качества: GenCode1 (5/5), GenCode2 (3/5), GenCode3 (4/5).')
add_paragraph(doc, 'Тесты "белого ящика" выявили скрытый дефект во второй реализации (отсутствие нормализации регистра). Обнаружена проблема излишней жесткости самих тестов при проверке локализованных исключений.')
add_paragraph(doc, 'Улучшение промпта: В промпт необходимо добавить требование поддержки любого регистра и строгого использования английского языка для генерации исключений.')

# Сериализация документа на диск
doc.save('Lab1_Report_Formatted.docx')
print("Документ Lab1_Report_Formatted.docx успешно сгенерирован!")
